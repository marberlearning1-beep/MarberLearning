from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x8B, 0x7A, 0xAB)
ACCENT2 = RGBColor(0x5B, 0xA8, 0x9A)
TEXT_DARK = RGBColor(0x2D, 0x2A, 0x3E)
TEXT_MUTED = RGBColor(0x6B, 0x68, 0x80)

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = TEXT_DARK


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(text, level=1, color=PRIMARY, size=None, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size or {1: 22, 2: 16, 3: 13}.get(level, 13))
    run.font.name = "Calibri"
    return p


def add_body(text, size=11, color=TEXT_DARK, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_callout(label, text, color=PRIMARY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F0EBF7")
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10.5)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(headers, rows, header_color="8B7AAB"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ---------- Cover ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
eyebrow = title.add_run("HANDOVER MANUAL\n")
eyebrow.font.size = Pt(10)
eyebrow.font.color.rgb = PRIMARY
eyebrow.bold = True
main = title.add_run("Website Handover Manual")
main.font.size = Pt(28)
main.bold = True
main.font.color.rgb = TEXT_DARK

add_body(
    "A plain-language reference for maintaining, understanding, and getting help with "
    "marberlearning.org — written for the Marber Learning Foundation team, not just developers.",
    size=12, color=TEXT_MUTED, space_after=14,
)

meta_table = doc.add_table(rows=4, cols=2)
meta_pairs = [
    ("Prepared for", "Roger Cortes — Marber Learning Foundation"),
    ("Prepared by", "Alesha Shahid — Website Developer"),
    ("Site", "marberlearning.org"),
    ("Last updated", "July 12, 2026"),
]
for i, (k, v) in enumerate(meta_pairs):
    meta_table.cell(i, 0).text = ""
    r = meta_table.cell(i, 0).paragraphs[0].add_run(k)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = TEXT_MUTED
    meta_table.cell(i, 1).text = ""
    r2 = meta_table.cell(i, 1).paragraphs[0].add_run(v)
    r2.font.size = Pt(10)
    r2.font.color.rgb = TEXT_DARK
meta_table.autofit = True

doc.add_paragraph().paragraph_format.space_after = Pt(6)
doc.add_page_break()

# ---------- 01 Overview ----------
add_heading("01 — Overview", level=2)
add_heading("What this document is", level=3, color=TEXT_DARK)
add_body(
    "This manual explains how the Marber Learning Foundation website works, in plain terms, "
    "so that anyone on your team — not just a developer — understands what exists, where it "
    "lives, and who to ask when something needs to change."
)
add_body(
    "A more technical companion set of notes (for a developer who picks up this project later) "
    "lives alongside the website's code as docs/DEVELOPER-NOTES.md and docs/DNS-SETUP.md. "
    "This manual is the front door; those files are the technical appendix."
)
add_table(
    ["Layer", "Built with", "In plain terms"],
    [
        ["Pages & navigation", "Next.js, React, TypeScript", "The framework that builds and displays each page"],
        ["Look & feel", "Tailwind CSS", "Controls colors, spacing, and layout"],
        ["Code storage", "GitHub", "Where every version of the site's code is saved"],
        ["Hosting", "Netlify", "Publishes the code as a live website"],
        ["Domain", "Namecheap", "Owns and points marberlearning.org at Netlify"],
    ],
)

# ---------- 02 Site map ----------
add_heading("02 — The four pages", level=2)
add_heading("What visitors can see today", level=3, color=TEXT_DARK)
add_body(
    "The site is built around four core pages, each with its own purpose in the visitor's "
    "journey from first impression to reaching out."
)
add_table(
    ["Page", "Purpose"],
    [
        ["1. Home", "The welcome page — introduces the foundation and links out to the other three."],
        ["2. About", "The foundation's mission, values, and story — where trust gets built."],
        ["3. Programs", "What tutoring programs are offered, and which schools/grade levels are served."],
        ["4. Contact", "A short form visitors use to reach the foundation directly by email."],
    ],
    header_color="5BA89A",
)

# ---------- 03 Making changes ----------
add_heading("03 — Making changes", level=2)
add_heading("How future edits happen", level=3, color=TEXT_DARK)
add_body(
    "This site is hand-built code, not a drag-and-drop editor — so text, photo, and layout "
    "changes are made by editing the code and publishing an update, rather than logging into "
    "a visual editor."
)
add_body(
    "The practical effect: any content change (new testimonial, updated program description, "
    "swapped photo) needs someone comfortable editing the code — either Alesha, on an hourly "
    "basis, or a future developer using the technical notes in docs/DEVELOPER-NOTES.md."
)
add_callout(
    "Good news",
    "the code was deliberately organized so that page text lives in clearly labeled, easy-to-find "
    "blocks near the top of each file — a developer should be able to make small text edits "
    "quickly rather than hunting through the whole site.",
)
add_heading("What happens after an edit is made", level=3, color=TEXT_DARK, space_before=6)
add_numbered([
    "A developer changes the relevant file and saves it to GitHub.",
    "Netlify notices the update automatically and rebuilds the site — usually within a minute or two.",
    "The change appears live at marberlearning.org with no further action needed.",
])

# ---------- 04 Contact form ----------
add_heading("04 — The contact form", level=2)
add_heading("How a visitor's message reaches you", level=3, color=TEXT_DARK)
add_body(
    "The Contact page form collects a name, email, phone, ZIP code, best time to call, and "
    "message. There is no database and no login involved — each submission is sent directly "
    "as an email, exactly as requested."
)
add_callout(
    "To confirm",
    "the original project brief named mcortes@marberlearnig.org as the recipient, while the "
    "current setup sends through info@marberlearning.org. Please confirm which inbox should "
    "receive form submissions so this can be finalized.",
    color=ACCENT2,
)

# ---------- 05 Accounts & access ----------
add_heading("05 — Accounts & access", level=2)
add_heading("Where the logins live", level=3, color=TEXT_DARK)
add_body(
    "At project handover, you should receive full login access to every account below. Use "
    "this table as a checklist — record the actual usernames and where the passwords are "
    "stored (ideally a password manager, not this document) once handover is complete."
)
add_table(
    ["Account", "Used for", "Username / notes"],
    [
        ["GitHub", "Stores the website's code and history", "to be recorded at handover"],
        ["Netlify", "Hosts and publishes the live site", "to be recorded at handover"],
        ["Namecheap", "Owns the marberlearning.org domain", "to be recorded at handover"],
        ["Zoho Mail", "Sends contact-form email notifications", "to be recorded at handover"],
    ],
)
add_callout(
    "Note",
    "The GitHub code currently exists under two accounts — one used during development, and "
    "one transferred to the foundation. Confirm with your developer which account Netlify is "
    "actively deploying from; that is the one that matters going forward.",
)

# ---------- 06 Domain & hosting ----------
add_heading("06 — Domain & hosting", level=2)
add_heading("How marberlearning.org stays online", level=3, color=TEXT_DARK)
add_body(
    "Namecheap owns the domain name itself and renews it yearly. A small set of settings there "
    "(called DNS records) point visitors from marberlearning.org to the website files hosted on "
    "Netlify. Netlify also automatically provides the padlock/secure-connection icon visitors "
    "see in their browser."
)
add_body(
    "The previous domain, marblelearning.org, now automatically forwards visitors to "
    "marberlearning.org, so old links and bookmarks continue to work."
)
add_body(
    "Full technical record of the exact DNS settings used lives in docs/DNS-SETUP.md alongside "
    "the website code."
)

# ---------- 07 Troubleshooting ----------
add_heading("07 — If something breaks", level=2)
add_heading("A quick first-response guide", level=3, color=TEXT_DARK)
add_bullets([
    "Site won't load at all: check Namecheap's DNS settings first — a record may have been changed or removed.",
    "Site loads but shows an old version: Netlify may still be building the latest update — this usually resolves within a few minutes.",
    "Browser shows “Not secure”: the security certificate may need renewing in Netlify's domain settings — this is usually automatic, but can occasionally need a manual refresh.",
    "Contact form stops sending emails: the Zoho email credentials in Netlify's environment settings may have expired or changed — check Site configuration → Environment variables.",
])
add_body(
    "For anything beyond these quick checks, the fastest path is contacting your developer "
    "directly rather than troubleshooting the code yourself."
)

# ---------- 08 Support ----------
add_heading("08 — Support", level=2)
add_heading("Getting help", level=3, color=TEXT_DARK)
add_body("For changes, questions, or issues beyond this manual, reach out directly:")
add_table(
    ["", ""],
    [
        ["Developer", "Alesha Shahid"],
        ["Email", "to be added"],
        ["Typical response", "Weekly Monday check-in, or by email between meetings"],
    ],
)

doc.add_paragraph().paragraph_format.space_before = Pt(12)
footer = add_body(
    "This manual is meant to travel with the project — keep it alongside "
    "docs/DEVELOPER-NOTES.md and docs/DNS-SETUP.md so that anyone who takes over the site in "
    "the future, whether that's you, Alesha, or someone new, can pick it up with no lost context.",
    size=9.5, color=TEXT_MUTED, italic=True,
)

doc.save(r"c:\Users\alesh\MarberTutoring\docs\Handover-Manual.docx")
print("saved")
